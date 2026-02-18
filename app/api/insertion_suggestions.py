from fastapi import APIRouter, HTTPException
from pydantic import BaseModel
from typing import List, Dict, Optional
import logging
import re

from app.services.gcp_llm_orchestrator import llm_orchestrator
from app.services.robust_document_modifier import robust_modifier

router = APIRouter()
logger = logging.getLogger(__name__)

class InsertionRequest(BaseModel):
    contract_id: str
    user_id: str
    clause_to_add: str
    context_hint: Optional[str] = None  # e.g., "after payment terms", "before termination"

class InsertionSuggestion(BaseModel):
    found_section: str
    insertion_point: str  # "after" or "before"
    context_preview: str  # Show surrounding text
    confidence: str  # "high", "medium", "low"

@router.post("/suggest")
async def suggest_insertion_point(request: InsertionRequest):
    """Suggest best location to insert new clause - hybrid AI + exact matching"""
    try:
        # Get document context
        contract_file = robust_modifier.find_contract_file(request.contract_id)
        if not contract_file:
            raise HTTPException(status_code=404, detail="Contract not found")
        
        # Load document and extract text
        from docx import Document
        doc = Document(contract_file)
        full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        
        found_section = None
        confidence = "low"
        reasoning = "Using default location"
        
        # PRIORITY 1: If user provided hint, use exact text matching
        if request.context_hint:
            import re
            hint_match = re.search(r'after\s+(.+)', request.context_hint, re.IGNORECASE)
            if hint_match:
                target_text = hint_match.group(1).strip()
                target_normalized = ' '.join(target_text.split())
                
                # Search for exact match in document
                for para in doc.paragraphs:
                    para_normalized = ' '.join(para.text.split())
                    if target_normalized.lower() in para_normalized.lower():
                        found_section = para.text.strip()
                        confidence = "high"
                        reasoning = f"Exact match found for: {target_text}"
                        break
        
        # PRIORITY 2: If no hint or no match found, use AI to suggest
        if not found_section:
            ai_prompt = f"""Analyze this contract and suggest where to insert the clause.

CONTRACT (first 3000 chars):
{full_text[:3000]}

CLAUSE TO ADD:
{request.clause_to_add}

Find a relevant section header or list where this clause fits semantically.
Return ONLY the exact text (5-15 words) from the contract where to insert.

For example, if adding documentation requirements, return "Copies of the following:" or "Required Documents:"

Return only the exact text, nothing else."""
            
            try:
                ai_response = await llm_orchestrator.chat_response(ai_prompt, [], [])
                
                # Extract suggested text
                suggested_text = ai_response['response'].strip().strip('"\'')
                
                # Verify AI suggestion exists in document
                suggested_normalized = ' '.join(suggested_text.split())
                for para in doc.paragraphs:
                    para_normalized = ' '.join(para.text.split())
                    if suggested_normalized.lower() in para_normalized.lower():
                        found_section = para.text.strip()
                        confidence = "medium"
                        reasoning = f"AI suggested: {suggested_text}"
                        break
            except Exception as e:
                logger.error(f"AI suggestion failed: {e}")
        
        # FALLBACK: Use default if nothing found
        if not found_section:
            found_section = "services"
            confidence = "low"
            reasoning = "No specific location found, using default"
        
        suggestion = {
            "found_section": found_section,
            "insertion_point": "after",
            "context_preview": f"Will insert after: {found_section}",
            "confidence": confidence,
            "reasoning": reasoning
        }
        
        return {
            "success": True,
            "contract_id": request.contract_id,
            "clause_to_add": request.clause_to_add,
            "suggestion": suggestion
        }
        
    except Exception as e:
        logger.error(f"Insertion suggestion failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

class ConfirmRequest(BaseModel):
    contract_id: str
    user_id: str
    clause_to_add: str
    found_section: str

@router.post("/confirm")
async def confirm_insertion(request: ConfirmRequest):
    """Apply the insertion after user confirms"""
    try:
        # Create modification with specific location
        modifications = [{
            "search_text": request.found_section,
            "replacement_text": request.clause_to_add,
            "action": "add",
            "explanation": f"Added new clause after {request.found_section}"
        }]
        
        # Apply modification
        result = robust_modifier.apply_modifications(request.contract_id, modifications)
        
        return {
            "success": True,
            "contract_id": request.contract_id,
            "changes_made": result['changes_made'],
            "download_urls": result['download_urls'],
            "message": f"Successfully added clause after {request.found_section}"
        }
        
    except Exception as e:
        logger.error(f"Insertion confirmation failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))
