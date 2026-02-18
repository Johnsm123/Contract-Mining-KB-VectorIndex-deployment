import os
import json
import logging
from pathlib import Path
from typing import Dict, List, Optional
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.shared import OxmlElement, qn
import re
from datetime import datetime

logger = logging.getLogger(__name__)

class RobustDocumentModifier:
    """Actually modify documents with real changes and color coding"""
    
    def __init__(self):
        self.storage_path = Path("local_storage")
        self.contracts_path = self.storage_path / "contracts"
        self.uploaded_path = self.storage_path / "uploaded_contracts"
        self.previews_path = self.storage_path / "previews"
        self.previews_path.mkdir(exist_ok=True)
    
    def find_contract_file(self, contract_id: str) -> Optional[Path]:
        """Find contract file by ID"""
        search_paths = [self.uploaded_path, self.contracts_path]
        
        for search_path in search_paths:
            if not search_path.exists():
                continue
                
            for file_path in search_path.glob("*.docx"):
                # Multiple search patterns
                if (f"_{contract_id}.docx" in file_path.name or 
                    f"_{contract_id}_" in file_path.name or
                    contract_id in file_path.name):
                    logger.info(f"Found contract file: {file_path}")
                    return file_path
        
        return None
    
    def get_document_context(self, contract_id: str) -> str:
        """Get document context from embeddings for better AI understanding"""
        try:
            embeddings_path = self.storage_path / "embeddings"
            
            for embedding_file in embeddings_path.glob(f"*{contract_id}*_embeddings.json"):
                with open(embedding_file, 'r', encoding='utf-8') as f:
                    embeddings_data = json.load(f)
                
                # Extract relevant text chunks
                context_text = "\n\n".join([chunk.get('text', '') for chunk in embeddings_data[:5]])
                return context_text[:3000]  # Limit context size
            
            return "No context available"
            
        except Exception as e:
            logger.warning(f"Could not load context: {e}")
            return "No context available"
    
    def apply_modifications(self, contract_id: str, modifications: List[Dict]) -> Dict:
        try:
            # Find the contract file
            contract_file = self.find_contract_file(contract_id)
            if not contract_file:
                available_files = []
                for path in [self.uploaded_path, self.contracts_path]:
                    if path.exists():
                        available_files.extend([f.name for f in path.glob("*.docx")])
                
                raise Exception(f"Contract {contract_id} not found. Available: {available_files[:3]}")
            
            # Load the document
            doc = Document(contract_file)
            changes_made = []
            
            logger.info(f"Processing {len(modifications)} modifications on {contract_file.name}")
            
            # Apply each modification
            for i, mod in enumerate(modifications):
                search_text = mod.get('search_text', '').strip()
                replacement_text = mod.get('replacement_text', '').strip()
                action = mod.get('action', 'replace')
                
                if not search_text:
                    continue
                
                logger.info(f"Modification {i+1}: {action} '{search_text}' -> '{replacement_text}'")
                
                # Track if modification was applied
                mod_applied = False
                
                # Process paragraphs
                for para_idx, paragraph in enumerate(doc.paragraphs):
                    # Normalize text for comparison (remove extra whitespace)
                    para_text_normalized = ' '.join(paragraph.text.split())
                    search_text_normalized = ' '.join(search_text.split())
                    
                    if search_text_normalized.lower() in para_text_normalized.lower():
                        original_text = paragraph.text
                        
                        if action == 'replace':
                            # Case-insensitive replace
                            import re
                            pattern = re.compile(re.escape(search_text), re.IGNORECASE)
                            new_text = pattern.sub(replacement_text, original_text)
                            paragraph.clear()
                            
                            # Add text with parts highlighted
                            parts = new_text.split(replacement_text)
                            for j, part in enumerate(parts):
                                if part:
                                    paragraph.add_run(part)
                                if j < len(parts) - 1:  # Not the last part
                                    highlighted_run = paragraph.add_run(replacement_text)
                                    highlighted_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                            
                            changes_made.append({
                                'location': f'paragraph_{para_idx}',
                                'original': original_text,
                                'modified': new_text,
                                'action': action
                            })
                            logger.info(f"Modified paragraph {para_idx} with YELLOW highlight")
                            
                        elif action == 'add':
                            # Add text as NEW PARAGRAPH after this one
                            from docx.oxml import OxmlElement
                            from docx.text.paragraph import Paragraph
                            
                            # Create new paragraph element
                            new_p_element = OxmlElement('w:p')
                            paragraph._element.addnext(new_p_element)
                            
                            # Create paragraph object from element
                            new_para = Paragraph(new_p_element, paragraph._parent)
                            
                            # Add the text with green highlighting
                            added_run = new_para.add_run(replacement_text)
                            added_run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                            
                            changes_made.append({
                                'location': f'new_paragraph_after_{para_idx}',
                                'original': '',
                                'modified': replacement_text,
                                'action': action
                            })
                            logger.info(f"Added as new paragraph after {para_idx} with GREEN highlight")
                            mod_applied = True
                            break
                            
                        elif action == 'delete':
                            # Find and strike through the text with case-insensitive search
                            import re
                            pattern = re.compile(re.escape(search_text), re.IGNORECASE)
                            
                            # Check if text exists
                            match = pattern.search(original_text)
                            if match:
                                paragraph.clear()
                                
                                # Get the text before and after the deleted portion
                                before_text = original_text[:match.start()]
                                after_text = original_text[match.end():]
                                deleted_text = match.group()
                                
                                # Add text before deletion
                                if before_text.strip():
                                    paragraph.add_run(before_text)
                                
                                # Add the deleted portion with strikethrough and red highlight
                                deleted_run = paragraph.add_run(deleted_text)
                                deleted_run.font.strike = True
                                deleted_run.font.highlight_color = WD_COLOR_INDEX.RED
                                
                                # Add text after deletion
                                if after_text.strip():
                                    paragraph.add_run(after_text)
                                
                                changes_made.append({
                                    'location': f'paragraph_{para_idx}',
                                    'original': original_text,
                                    'modified': f'{before_text}{deleted_text}{after_text}',
                                    'action': action
                                })
                                logger.info(f"Deleted from paragraph {para_idx} with RED strikethrough")
                
                # If it's an 'add' action without search_text match, add to the first suitable paragraph
                if action == 'add' and not mod_applied:
                    # Find a paragraph that contains common contract terms
                    suitable_paragraphs = []
                    for para_idx, paragraph in enumerate(doc.paragraphs):
                        if any(term in paragraph.text.lower() for term in ['service', 'agreement', 'contract', 'terms', 'conditions', 'requirements']):
                            suitable_paragraphs.append((para_idx, paragraph))
                    
                    if suitable_paragraphs:
                        # Use the first suitable paragraph
                        para_idx, paragraph = suitable_paragraphs[0]
                        original_text = paragraph.text
                        
                        # Add text with green highlighting at the END of paragraph
                        added_run = paragraph.add_run(' ' + replacement_text)
                        added_run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                        
                        changes_made.append({
                            'location': f'paragraph_{para_idx}',
                            'original': original_text,
                            'modified': original_text + ' ' + replacement_text,
                            'action': action
                        })
                        mod_applied = True
                        logger.info(f"Added clause to paragraph {para_idx} with GREEN highlight")
                        break
                    
                    if not mod_applied and not suitable_paragraphs:
                        # If still no suitable paragraph, add to the end of the document
                        new_paragraph = doc.add_paragraph()
                        added_run = new_paragraph.add_run(replacement_text)
                        added_run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                        
                        changes_made.append({
                            'location': 'new_paragraph_end',
                            'original': '',
                            'modified': replacement_text,
                            'action': action
                        })
                        mod_applied = True
                        logger.info(f"Added clause as new paragraph at end with GREEN highlight")
                
                # Process tables only if not already applied
                for table_idx, table in enumerate(doc.tables):
                    if mod_applied and action == 'add':
                        break
                    for row_idx, row in enumerate(table.rows):
                        for cell_idx, cell in enumerate(row.cells):
                            if action == 'replace':
                                # Process each paragraph in the cell
                                for para in cell.paragraphs:
                                    if search_text.lower() in para.text.lower():
                                        original_text = para.text
                                        # Case-insensitive replace
                                        import re
                                        pattern = re.compile(re.escape(search_text), re.IGNORECASE)
                                        new_text = pattern.sub(replacement_text, original_text)
                                        
                                        # Clear and rebuild paragraph
                                        para.clear()
                                        
                                        # Add text with only the replacement highlighted
                                        parts = new_text.split(replacement_text)
                                        for j, part in enumerate(parts):
                                            if part:
                                                para.add_run(part)
                                            if j < len(parts) - 1:
                                                highlighted_run = para.add_run(replacement_text)
                                                highlighted_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                                        
                                        changes_made.append({
                                            'location': f'table_{table_idx}_row_{row_idx}_cell_{cell_idx}',
                                            'original': original_text,
                                            'modified': new_text,
                                            'action': action
                                        })
                                        logger.info(f"Modified table cell: '{original_text}' -> '{new_text}'")
                                    
                            elif action == 'add':
                                for para in cell.paragraphs:
                                    # Normalize text for comparison
                                    para_text_normalized = ' '.join(para.text.split())
                                    search_text_normalized = ' '.join(search_text.split())
                                    
                                    if search_text_normalized.lower() in para_text_normalized.lower():
                                        original_text = para.text
                                        added_run = para.add_run(' ' + replacement_text)
                                        added_run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                                        
                                        changes_made.append({
                                            'location': f'table_{table_idx}_row_{row_idx}_cell_{cell_idx}',
                                            'original': original_text,
                                            'modified': original_text + ' ' + replacement_text,
                                            'action': action
                                        })
                                        logger.info(f"Added to table cell with GREEN highlight")
                                        mod_applied = True
                                        break
                                        
                            elif action == 'delete':
                                for para in cell.paragraphs:
                                    # Normalize text for comparison
                                    para_text_normalized = ' '.join(para.text.split())
                                    search_text_normalized = ' '.join(search_text.split())
                                    
                                    if search_text_normalized.lower() in para_text_normalized.lower():
                                        original_text = para.text
                                        import re
                                        pattern = re.compile(re.escape(search_text), re.IGNORECASE)
                                        
                                        match = pattern.search(original_text)
                                        if match:
                                            para.clear()
                                            
                                            # Get the text before and after the deleted portion
                                            before_text = original_text[:match.start()]
                                            after_text = original_text[match.end():]
                                            deleted_text = match.group()
                                            
                                            # Add text before deletion
                                            if before_text.strip():
                                                para.add_run(before_text)
                                            
                                            # Add the deleted portion with strikethrough and red highlight
                                            deleted_run = para.add_run(deleted_text)
                                            deleted_run.font.strike = True
                                            deleted_run.font.highlight_color = WD_COLOR_INDEX.RED
                                            
                                            # Add text after deletion
                                            if after_text.strip():
                                                para.add_run(after_text)
                                            
                                            changes_made.append({
                                                'location': f'table_{table_idx}_row_{row_idx}_cell_{cell_idx}',
                                                'original': original_text,
                                                'modified': f'{before_text}{deleted_text}{after_text}',
                                                'action': action
                                            })
                                            logger.info(f"Deleted from table cell with RED strikethrough")
                
                if not changes_made:
                    logger.warning(f"No matches found for: '{search_text}'")
            
            # Save the modified document
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_filename = f"modified_{contract_id}_{timestamp}.docx"
            output_path = self.previews_path / output_filename
            
            doc.save(output_path)
            logger.info(f"Saved modified document: {output_path}")
            
            # Generate summary
            return {
                'success': True,
                'contract_file': str(contract_file),
                'output_file': str(output_path),
                'changes_made': len(changes_made),
                'changes_detail': changes_made,
                'download_urls': {
                    'word': f'/api/modify/download/{output_filename}'
                }
            }
            
        except Exception as e:
            logger.error(f"Document modification failed: {e}")
            raise e

# Global instance
robust_modifier = RobustDocumentModifier()