"""Amendment generation with color-coded changes"""
import logging
from typing import Dict, List
from datetime import datetime
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX
from google.cloud import storage
import json
import os

from app.config.settings import settings
from app.services.gcp_document_processor import document_processor

logger = logging.getLogger(__name__)

class AmendmentGenerator:
    """Generate color-coded amendments"""
    
    def __init__(self):
        self.storage_client = storage.Client(project=settings.gcp_project_id)
        self.amendments_bucket = self.storage_client.bucket(settings.gcs_amendments_bucket)
        self.temps_bucket = self.storage_client.bucket(settings.gcs_embeddings_bucket)
    
    async def generate_amendment(self, contract_id: str, user_id: str, modifications: List[Dict]) -> Dict:
        """Generate color-coded amendment document"""
        try:
            # Load original contract
            contract_path = os.path.join("local_storage", "contracts", f"{contract_id}.docx")
            
            if not os.path.exists(contract_path):
                raise FileNotFoundError(f"Contract {contract_id} not found")
            
            doc = Document(contract_path)
            
            # Track changes
            changes_detail = []
            changes_made = 0
            
            # Apply each modification with color coding
            for mod in modifications:
                action = mod.get('action', 'replace')
                search_text = mod.get('search_text', '')
                new_text = mod.get('new_text', '')
                explanation = mod.get('explanation', '')
                
                if action == 'add':
                    # GREEN for additions
                    result = self._add_text(doc, search_text, new_text)
                    if result:
                        changes_made += 1
                        changes_detail.append({
                            'action': 'add',
                            'location': f"After '{search_text}'",
                            'text': new_text,
                            'color': 'green',
                            'explanation': explanation
                        })
                
                elif action == 'delete':
                    # RED strikethrough for deletions
                    result = self._delete_text(doc, search_text)
                    if result:
                        changes_made += 1
                        changes_detail.append({
                            'action': 'delete',
                            'text': search_text,
                            'color': 'red',
                            'explanation': explanation
                        })
                
                elif action == 'replace':
                    # YELLOW for replacements
                    result = self._replace_text(doc, search_text, new_text)
                    if result:
                        changes_made += 1
                        changes_detail.append({
                            'action': 'replace',
                            'old_text': search_text,
                            'new_text': new_text,
                            'color': 'yellow',
                            'explanation': explanation
                        })
            
            # Generate amendment ID
            amendment_id = f"amendment_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            
            # Save locally
            preview_dir = os.path.join("local_storage", "previews")
            os.makedirs(preview_dir, exist_ok=True)
            local_path = os.path.join(preview_dir, f"{amendment_id}.docx")
            doc.save(local_path)
            
            # Upload to GCS amendments bucket
            blob = self.amendments_bucket.blob(f"{amendment_id}.docx")
            blob.upload_from_filename(local_path)
            gcs_amendment_path = f"gs://{settings.gcs_amendments_bucket}/{amendment_id}.docx"
            
            # Generate embeddings for the amendment
            embeddings_path = await self._generate_embeddings(amendment_id, local_path)
            
            logger.info(f"Amendment generated: {amendment_id} with {changes_made} changes")
            
            return {
                'success': True,
                'amendment_id': amendment_id,
                'contract_id': contract_id,
                'changes_made': changes_made,
                'changes_detail': changes_detail,
                'download_urls': {
                    'docx': f"/api/modify/download/{amendment_id}.docx"
                },
                'gcs_paths': {
                    'amendment': gcs_amendment_path,
                    'embeddings': embeddings_path
                },
                'message': f"Amendment generated with {changes_made} color-coded changes"
            }
            
        except Exception as e:
            logger.error(f"Amendment generation failed: {e}")
            raise
    
    def _add_text(self, doc: Document, search_text: str, new_text: str) -> bool:
        """Add text with GREEN highlight"""
        try:
            for paragraph in doc.paragraphs:
                if search_text.lower() in paragraph.text.lower():
                    # Add new paragraph after this one with green highlight
                    new_para = paragraph.insert_paragraph_before("")
                    run = new_para.add_run(f"\n{new_text}")
                    run.font.color.rgb = RGBColor(0, 128, 0)  # Green
                    run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                    run.font.size = Pt(11)
                    return True
            return False
        except Exception as e:
            logger.error(f"Add text failed: {e}")
            return False
    
    def _delete_text(self, doc: Document, search_text: str) -> bool:
        """Mark text as deleted with RED strikethrough"""
        try:
            for paragraph in doc.paragraphs:
                if search_text in paragraph.text:
                    # Find and mark the text
                    for run in paragraph.runs:
                        if search_text in run.text:
                            run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                            run.font.strike = True
                            return True
            return False
        except Exception as e:
            logger.error(f"Delete text failed: {e}")
            return False
    
    def _replace_text(self, doc: Document, old_text: str, new_text: str) -> bool:
        """Replace text with YELLOW highlight"""
        try:
            for paragraph in doc.paragraphs:
                if old_text in paragraph.text:
                    # Mark old text as strikethrough
                    for run in paragraph.runs:
                        if old_text in run.text:
                            run.font.strike = True
                            run.font.color.rgb = RGBColor(128, 128, 128)  # Gray
                    
                    # Add new text with yellow highlight
                    new_run = paragraph.add_run(f" {new_text}")
                    new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    new_run.font.color.rgb = RGBColor(0, 0, 0)  # Black
                    new_run.font.size = Pt(11)
                    return True
            return False
        except Exception as e:
            logger.error(f"Replace text failed: {e}")
            return False
    
    async def _generate_embeddings(self, amendment_id: str, file_path: str) -> str:
        """Generate embeddings for amendment and store in temps/amendments/"""
        try:
            # Extract text from document
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            # Generate embeddings using document processor
            chunks = self._split_text(text)
            embeddings_data = []
            
            for i, chunk in enumerate(chunks):
                embedding = await document_processor._generate_embedding(chunk)
                embeddings_data.append({
                    'chunk_id': i,
                    'text': chunk,
                    'embedding': embedding,
                    'amendment_id': amendment_id
                })
            
            # Store in temps bucket under kb/amendments/
            embeddings_json = json.dumps(embeddings_data, indent=2)
            blob_path = f"kb/amendments/{amendment_id}_embeddings.json"
            blob = self.temps_bucket.blob(blob_path)
            blob.upload_from_string(embeddings_json, content_type='application/json')
            
            gcs_path = f"gs://{settings.gcs_embeddings_bucket}/{blob_path}"
            logger.info(f"Amendment embeddings stored: {gcs_path}")
            return gcs_path
            
        except Exception as e:
            logger.error(f"Embedding generation failed: {e}")
            return ""
    
    def _split_text(self, text: str, chunk_size: int = 500) -> List[str]:
        """Split text into chunks"""
        words = text.split()
        chunks = []
        current_chunk = []
        current_size = 0
        
        for word in words:
            current_chunk.append(word)
            current_size += len(word) + 1
            
            if current_size >= chunk_size:
                chunks.append(" ".join(current_chunk))
                current_chunk = []
                current_size = 0
        
        if current_chunk:
            chunks.append(" ".join(current_chunk))
        
        return chunks

amendment_generator = AmendmentGenerator()
